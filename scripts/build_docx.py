#!/usr/bin/env python3
"""
Build the monthly newsletter .docx from a content JSON file.

Usage: python scripts/build_docx.py YYYYMM

Reads:  reports/YYYYMM_content.json
Writes: reports/YYYY MM Newsletter.docx
"""

import sys
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).parent.parent

DUNE_FINANCIALS = "https://dune.com/nexus_mutual/capital-pool-and-ownership"
DUNE_RAMM = "https://dune.com/nexus_mutual/ramm"
DUNE_COVERS = "https://dune.com/nexus_mutual/covers"

MONTH_NAMES = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


def add_hyperlink(paragraph, text: str, url: str):
    """Append a blue underlined hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_section_header(doc: Document, text: str):
    """Add a bold normal-style paragraph used as a section header."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def add_chart(doc: Document, image_path: str):
    """Add a chart image, centred, at a fixed width."""
    p = doc.add_paragraph()
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(6))
    return p


def build_newsletter(content: dict, output_path: Path) -> None:
    doc = Document()

    month_year = content["month_year"]          # e.g. "October 2025"
    month_abbr = month_year[:3]                 # e.g. "Oct"
    year_short = month_year.split()[1][2:]      # e.g. "25"
    table = content["investment_table"]

    # ── Title ────────────────────────────────────────────────────────────────
    doc.add_heading(f"Investment Committee Newsletter - {month_year}", 1)

    # ── Intro ─────────────────────────────────────────────────────────────────
    doc.add_paragraph(content["intro"])

    # ── State of the Capital Pool ─────────────────────────────────────────────
    add_section_header(doc, "State of the Capital Pool")

    add_section_header(doc, "Monthly Change - ETH value")
    doc.add_paragraph(content["capital_pool_eth_paragraph"])
    doc.add_paragraph(
        "The various impacts on the capital pool are summarised in the waterfall chart below."
    )
    add_chart(doc, content["chart_paths"]["waterfall_eth"])
    doc.add_paragraph(
        "The cover fee income is net of distribution commissions and excludes covers paid for in NXM. "
        "In such a case, the cover fee gets burned and there is no change in the Capital Pool."
    )

    add_section_header(doc, "Monthly Change in NXM Book Value")
    doc.add_paragraph(content["capital_pool_bv_paragraph"])
    doc.add_paragraph(
        "The various impacts on the capital pool are summarised in the waterfall chart below."
    )
    add_chart(doc, content["chart_paths"]["waterfall_bv"])

    p = doc.add_paragraph("\u2192 Members can track protocol\u2019s revenue on the ")
    add_hyperlink(p, "Financials Dune Dashboard", DUNE_FINANCIALS)
    p = doc.add_paragraph("\u2192 Members can track in/outflows on the ")
    add_hyperlink(p, "Ratcheting AMM Dune Dashboard", DUNE_RAMM)
    p = doc.add_paragraph("\u2192 Members can track the cover income on the ")
    add_hyperlink(p, "Covers Dune Dashboard", DUNE_COVERS)

    # ── End of Month Pool Split ───────────────────────────────────────────────
    add_section_header(doc, "End of Month Pool Split")
    doc.add_paragraph(
        f"The split of the Capital Pool at the end of {month_abbr} '{year_short} in ETH terms is as follows."
    )
    add_chart(doc, content["chart_paths"]["pie_chart"])
    p = doc.add_paragraph("\u2192 Members can find the updated split at any time on the ")
    add_hyperlink(p, "Capital Pool and Ownership Dune Dashboard", DUNE_FINANCIALS)

    # ── State of the Investments ──────────────────────────────────────────────
    add_section_header(doc, "State of the Investments")
    doc.add_paragraph(
        f"In the last month, the Mutual earned {table['total_return']:.1f} ETH on its investments, "
        "overall, as broken down below."
    )

    doc.add_paragraph(f"stETH Monthly Return: {table['steth_return']}")
    doc.add_paragraph(f"stETH Monthly APY: {table['steth_apy']:.3f}%")
    doc.add_paragraph(f"rETH Monthly Return: {table['reth_return']}")
    doc.add_paragraph(f"rETH Monthly APY: {table['reth_apy']:.3f}%")
    doc.add_paragraph(f"Enzyme Vault Monthly Return: {table['enzyme_return']}")
    doc.add_paragraph(f"Enzyme Vault Monthly APY: {table['enzyme_apy']:.3f}%")
    doc.add_paragraph("Enzyme Vault includes EtherFi investments")
    doc.add_paragraph(f"Total ETH Earned: {table['total_return']}")
    doc.add_paragraph(f"Total Monthly APY: {table['total_apy']:.3f}%")
    doc.add_paragraph("Based on average Capital Pool amount over the monthly period")
    doc.add_paragraph("All returns after fees")

    add_chart(doc, content["chart_paths"]["investment_returns"])

    doc.add_paragraph(content["investment_summary_paragraph"])

    doc.save(output_path)
    print(f"Newsletter saved: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python build_docx.py YYYYMM", file=sys.stderr)
        sys.exit(1)

    yyyymm = sys.argv[1]
    content_path = REPO_ROOT / "reports" / f"{yyyymm}_content.json"

    if not content_path.exists():
        print(f"Error: content file not found: {content_path}", file=sys.stderr)
        sys.exit(1)

    with open(content_path) as f:
        content = json.load(f)

    year = yyyymm[:4]
    month = int(yyyymm[4:])
    output_path = REPO_ROOT / "reports" / f"{year} {month:02d} Newsletter.docx"

    build_newsletter(content, output_path)


if __name__ == "__main__":
    main()
